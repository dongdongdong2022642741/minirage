"""Tests for the docparser module. Run with:  python -m unittest discover tests"""

import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docparser import Document, load_documents, parse_file

DOCS_DIR = Path(__file__).resolve().parent.parent / "docs"


class ParseFileTests(unittest.TestCase):
    def test_parses_markdown(self):
        doc = parse_file(DOCS_DIR / "sample_01.md")
        self.assertIsInstance(doc, Document)
        self.assertEqual(doc.filename, "sample_01.md")
        self.assertEqual(doc.suffix, "md")
        self.assertIn("RAG", doc.text)
        self.assertGreater(doc.char_count, 0)
        self.assertTrue(doc.doc_id)

    def test_parses_txt(self):
        doc = parse_file(DOCS_DIR / "sample_02.txt")
        self.assertEqual(doc.filename, "sample_02.txt")
        self.assertEqual(doc.suffix, "txt")

    def test_parses_gbk_file_via_fallback(self):
        doc = parse_file(DOCS_DIR / "sample_03.txt")
        self.assertIn("GBK 编码示例", doc.text)
        self.assertEqual(doc.encoding, "gb18030")

    def test_missing_file_raises(self):
        with self.assertRaises(FileNotFoundError):
            parse_file(DOCS_DIR / "nope.md")

    def test_unsupported_suffix_raises(self):
        with tempfile.TemporaryDirectory() as tmp:
            bad = Path(tmp) / "data.bin"
            bad.write_bytes(b"\x00\x01")
            with self.assertRaises(ValueError):
                parse_file(bad)

    def test_empty_file_raises(self):
        with tempfile.TemporaryDirectory() as tmp:
            empty = Path(tmp) / "empty.md"
            empty.write_text("", encoding="utf-8")
            with self.assertRaises(ValueError):
                parse_file(empty)

    def test_parses_docx_headings_paragraphs_and_tables(self):
        from docx import Document as DocxDocument

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "manual.docx"
            source = DocxDocument()
            source.add_heading("操作手册", level=1)
            source.add_paragraph("启动设备前检查电源。")
            table = source.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "项目"
            table.cell(0, 1).text = "要求"
            table.cell(1, 0).text = "温度"
            table.cell(1, 1).text = "20 度"
            source.save(path)

            doc = parse_file(path)
            self.assertEqual(doc.suffix, "docx")
            self.assertEqual(doc.encoding, "docx")
            self.assertIn("# 操作手册", doc.text)
            self.assertIn("| 项目 | 要求 |", doc.text)
            self.assertIn("| 温度 | 20 度 |", doc.text)

    def test_parses_html_as_markdown_without_scripts(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "manual.html"
            path.write_text(
                "<html><body><h1>操作手册</h1><p>检查电源。</p>"
                "<script>secret()</script><table><tr><th>项目</th><th>要求</th>"
                "</tr><tr><td>温度</td><td>20 度</td></tr></table></body></html>",
                encoding="utf-8",
            )

            doc = parse_file(path)
            self.assertEqual(doc.suffix, "html")
            self.assertIn("# 操作手册", doc.text)
            self.assertIn("| 温度 | 20 度 |", doc.text)
            self.assertNotIn("secret", doc.text)

    def test_parses_pdf_pages_with_page_headings(self):
        class FakePage:
            def __init__(self, text):
                self.text = text

            def extract_text(self):
                return self.text

        class FakeReader:
            def __init__(self, _path):
                self.pages = [FakePage("第一页内容"), FakePage("第二页内容")]

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "manual.pdf"
            path.write_bytes(b"%PDF-test")
            with patch("pypdf.PdfReader", FakeReader):
                doc = parse_file(path)

            self.assertEqual(doc.encoding, "pdf")
            self.assertIn("## 第 1 页", doc.text)
            self.assertIn("## 第 2 页", doc.text)
            self.assertIn("第二页内容", doc.text)

    def test_scanned_pdf_without_text_fails_when_ocr_disabled(self):
        class FakePage:
            def extract_text(self):
                return ""

        class FakeReader:
            def __init__(self, _path):
                self.pages = [FakePage()]

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "scan.pdf"
            path.write_bytes(b"%PDF-test")
            with patch("pypdf.PdfReader", FakeReader), \
                    patch("docparser.ocr.ocr_enabled", return_value=False):
                with self.assertRaisesRegex(ValueError, "OCR"):
                    parse_file(path)

    def test_scanned_pdf_uses_ocr_when_enabled(self):
        class FakePage:
            def extract_text(self):
                return ""

        class FakeReader:
            def __init__(self, _path):
                self.pages = [FakePage()]

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "scan.pdf"
            path.write_bytes(b"%PDF-test")
            with patch("pypdf.PdfReader", FakeReader), \
                    patch("docparser.ocr.ocr_enabled", return_value=True), \
                    patch("docparser.ocr.ocr_pdf_page",
                          return_value="OCR 识别的中文内容") as fake_ocr:
                doc = parse_file(path)

            self.assertIn("## 第 1 页", doc.text)
            self.assertIn("OCR 识别的中文内容", doc.text)
            fake_ocr.assert_called_once()

    def test_ocr_failure_on_one_page_does_not_kill_text_pages(self):
        class FakeTextPage:
            def extract_text(self):
                return "第一页有文本层"

        class FakeEmptyPage:
            def extract_text(self):
                return ""

        class FakeReader:
            def __init__(self, _path):
                self.pages = [FakeTextPage(), FakeEmptyPage()]

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "mixed.pdf"
            path.write_bytes(b"%PDF-test")
            with patch("pypdf.PdfReader", FakeReader), \
                    patch("docparser.ocr.ocr_enabled", return_value=True), \
                    patch("docparser.ocr.ocr_pdf_page",
                          side_effect=RuntimeError("engine boom")):
                doc = parse_file(path)

            self.assertIn("## 第 1 页", doc.text)
            self.assertIn("第一页有文本层", doc.text)
            self.assertNotIn("## 第 2 页", doc.text)


class LoadDocumentsTests(unittest.TestCase):
    def test_loads_md_and_txt(self):
        result = load_documents(DOCS_DIR)
        docs = {d.filename: d for d in result.documents}
        self.assertIn("sample_01.md", docs)
        self.assertIn("sample_02.txt", docs)
        self.assertIn("sample_03.txt", docs)

    def test_empty_file_recorded_as_issue(self):
        result = load_documents(DOCS_DIR)
        reasons = {i.reason for i in result.issues}
        self.assertIn("empty_file", reasons)

    def test_undecodable_file_recorded_as_issue(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            (root / "broken.txt").write_bytes(b"\x81")
            result = load_documents(root)
            self.assertEqual(result.documents, [])
            self.assertEqual(len(result.issues), 1)
            self.assertEqual(result.issues[0].reason, "decode_error")

    def test_doc_id_is_stable_across_calls(self):
        first = load_documents(DOCS_DIR).documents
        second = load_documents(DOCS_DIR).documents
        first_ids = {d.filename: d.doc_id for d in first}
        second_ids = {d.filename: d.doc_id for d in second}
        self.assertEqual(first_ids, second_ids)

    def test_doc_id_differs_between_documents(self):
        result = load_documents(DOCS_DIR)
        ids = [d.doc_id for d in result.documents]
        self.assertEqual(len(ids), len(set(ids)))

    def test_missing_directory_raises(self):
        with self.assertRaises(FileNotFoundError):
            load_documents(Path(tempfile.gettempdir()) / "definitely-not-here")

    def test_ignores_non_doc_files_and_dirs(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            (root / "notes.md").write_text("hello", encoding="utf-8")
            (root / "image.png").write_bytes(b"\x89PNG")
            (root / "sub").mkdir()
            (root / "sub" / "inner.md").write_text("inner", encoding="utf-8")
            result = load_documents(root)
            names = [d.filename for d in result.documents]
            self.assertEqual(names, ["notes.md"])

    def test_binary_parse_failure_is_recorded(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            (root / "broken.docx").write_bytes(b"not-a-docx")
            result = load_documents(root)
            self.assertEqual(result.documents, [])
            self.assertEqual(result.issues[0].reason, "parse_error")


class DocumentShapeTests(unittest.TestCase):
    def test_required_fields_present(self):
        doc = parse_file(DOCS_DIR / "sample_01.md")
        self.assertTrue(hasattr(doc, "doc_id"))
        self.assertTrue(hasattr(doc, "filename"))
        self.assertTrue(hasattr(doc, "text"))

    def test_frozen_immutability(self):
        doc = parse_file(DOCS_DIR / "sample_01.md")
        with self.assertRaises(Exception):
            doc.text = "changed"  # type: ignore[misc]


if __name__ == "__main__":
    unittest.main(verbosity=2)
